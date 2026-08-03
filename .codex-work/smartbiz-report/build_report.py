from pathlib import Path
from copy import deepcopy
import shutil
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path.cwd()
REFERENCE = Path(r'C:\Users\sa\Downloads\FINAL-YEAR-PROJECT-REPORT.docx')
OUT = ROOT / 'deliverables' / 'SmartBiz-Final-Year-Project-Report.docx'
WORK = ROOT / '.codex-work' / 'smartbiz-report'
DIAGRAMS = ROOT / 'deliverables' / 'diagrams'
OUT.parent.mkdir(parents=True, exist_ok=True)
WORK.mkdir(parents=True, exist_ok=True)

shutil.copy2(REFERENCE, OUT)

with zipfile.ZipFile(REFERENCE) as z:
    media = [n for n in z.namelist() if n.startswith('word/media/')]
    logo_name = next((n for n in media if n.lower().endswith(('image1.jpeg','image1.jpg','image1.png'))), media[0])
    logo = WORK / Path(logo_name).name
    logo.write_bytes(z.read(logo_name))

doc = Document(OUT)

# Clear body content while retaining the final section properties and the reference package assets.
body = doc._element.body
sect_pr = body.sectPr
for child in list(body):
    if child is not sect_pr:
        body.remove(child)

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_repeat_header(row):
    set_repeat_table_header(row)

def set_outline_level(style, level):
    p_pr = style.element.get_or_add_pPr()
    node = p_pr.find(qn('w:outlineLvl'))
    if node is None:
        node = OxmlElement('w:outlineLvl')
        p_pr.append(node)
    node.set(qn('w:val'), str(level))

def ensure_style(name, size, bold=False, italic=False, align=None, before=0, after=6,
                 line=1.0, keep=False, outline=None):
    styles = doc.styles
    if name in styles:
        st = styles[name]
    else:
        st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = styles['Normal']
    st.font.name = 'Arial'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.italic = italic
    pf = st.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep
    if align is not None:
        pf.alignment = align
    if outline is not None:
        set_outline_level(st, outline)
    return st

# Reconstructed styles based on the reference audit.
normal = doc.styles['Normal']
normal.font.name = 'Arial'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
normal.font.size = Pt(11)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)

ensure_style('SB Front Heading', 14, True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18, keep=True, outline=0)
ensure_style('SB Chapter', 14, True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18, keep=True, outline=0)
ensure_style('SB Heading 2', 12, True, before=12, after=6, keep=True, outline=1)
ensure_style('SB Heading 3', 11, True, before=9, after=4, keep=True, outline=2)
ensure_style('SB Heading 4', 11, True, italic=True, before=6, after=3, keep=True)
ensure_style('SB Figure Caption', 9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=9)
ensure_style('SB Table Caption', 9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=7, after=4)
ensure_style('SB TOC Field', 11, align=WD_ALIGN_PARAGRAPH.LEFT, after=0, line=1.15)
ensure_style('SB Reference', 10, align=WD_ALIGN_PARAGRAPH.LEFT, after=5, line=1.15)
ensure_style('SB Compact', 10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3, line=1.15)

def configure_section(section):
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    section.different_first_page_header_footer = False

def set_page_numbering(section, fmt=None, start=None):
    sect = section._sectPr
    old = sect.find(qn('w:pgNumType'))
    if old is not None:
        sect.remove(old)
    pg = OxmlElement('w:pgNumType')
    if fmt:
        pg.set(qn('w:fmt'), fmt)
    if start is not None:
        pg.set(qn('w:start'), str(start))
    sect.append(pg)

def add_field(paragraph, instruction, placeholder='Right-click and update field'):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    begin.set(qn('w:dirty'), 'true')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instruction
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = placeholder
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, text, end])

def set_footer_page(section):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, 'PAGE', '1')
    for r in p.runs:
        r.font.name = 'Arial'; r.font.size = Pt(10)

def blank_footer(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.clear()

def add_page_break():
    doc.add_page_break()

def add_heading(text, level=2):
    style = {1:'SB Chapter', 2:'SB Heading 2', 3:'SB Heading 3', 4:'SB Heading 4'}[level]
    p = doc.add_paragraph(text, style=style)
    if level == 1:
        p.paragraph_format.page_break_before = True
    return p

def add_front_heading(text):
    return doc.add_paragraph(text.upper(), style='SB Front Heading')

def add_body(text, bold_lead=None, style=None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p

def add_body_block(block):
    for text in [x.strip() for x in block.strip().split('\n\n') if x.strip()]:
        add_body(text)

def add_bullets(items, numbered=False):
    for index, item in enumerate(items, 1):
        p = doc.add_paragraph(style='SB Compact')
        p.paragraph_format.left_indent = Inches(.25)
        p.paragraph_format.first_line_indent = Inches(-.18)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.3
        p.add_run(f'{index}. ' if numbered else '• ')
        p.add_run(item)

def add_caption(text, kind='Figure'):
    return doc.add_paragraph(text, style='SB Figure Caption' if kind == 'Figure' else 'SB Table Caption')

def add_figure(filename, caption, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(DIAGRAMS / filename), width=Inches(width))
    add_caption(caption, 'Figure')

def add_table(caption, headers, rows, widths=None, font_size=9.5):
    add_caption(caption, 'Table')
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    tbl.autofit = False
    hdr = tbl.rows[0]
    set_repeat_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, 'E7E6E6')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text)); r.bold = True; r.font.name='Arial'; r.font.size=Pt(font_size)
        if widths: cell.width = Inches(widths[i])
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value)); r.font.name='Arial'; r.font.size=Pt(font_size)
            if widths: cell.width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return tbl

def add_code_block(text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0,0)
    set_cell_shading(cell, 'F5F5F5')
    set_cell_margins(cell, 120, 140, 120, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = 'Courier New'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Courier New'); r.font.size = Pt(8.5)

for s in doc.sections:
    configure_section(s)
    blank_footer(s)

# Cover (section 1)
s1 = doc.sections[0]
blank_footer(s1)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
p.add_run().add_picture(str(logo), width=Inches(1.35))
for text, size, bold, before, after in [
    ('FINAL YEAR PROJECT REPORT',14,True,6,26),
    ('SMARTBIZ',14,True,0,6),
    ('A Mobile-First Business Management System for Small Businesses in Nepal',12,True,0,30),
    ('Prepared and submitted by',11,False,0,10),
    ('SANSKAR KAYASTHA ([STUDENT ID])',12,True,0,8),
    ('8th Semester',11,True,0,28),
    ('In partial fulfilment of the requirements for the degree of',11,False,0,10),
    ('BACHELOR OF INFORMATION AND COMMUNICATION TECHNOLOGY',11,True,0,18),
    ('School of Science and Technology',11,True,0,4),
    ('Virinchi College',11,True,0,4),
    ('Asia e University',11,True,0,26),
    ('JULY 2026',11,True,0,0),
]:
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1
    r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(size); r.bold=bold

# Section 2: declaration and approval, Roman numbering.
s2 = doc.add_section(WD_SECTION.NEW_PAGE); configure_section(s2); set_page_numbering(s2,'lowerRoman',1); set_footer_page(s2)
add_front_heading('Declaration')
add_body_block('''I hereby declare that the project entitled “SmartBiz: A Mobile-First Business Management System for Small Businesses in Nepal” is my original work carried out in partial fulfilment of the requirements for the Bachelor of Information and Communication Technology degree. The report has not been submitted, in whole or in part, for any other academic award.

All sources of information used in this work have been acknowledged in the reference section. The design, implementation, testing evidence and analysis presented here are based on the SmartBiz repository and the development work completed during the project period. Where third-party frameworks, hosted services and documentation were used, their contribution has been identified appropriately.''')
doc.add_paragraph('\n\n')
add_table('Declaration signature record', ['Candidate','Signature','Date'], [['Sanskar Kayastha','____________________','____________________']], [2.4,2.0,1.8],10)
add_page_break()
add_front_heading('Approved By')
add_body('This final year project report has been examined and approved as meeting the academic requirements of the Bachelor of Information and Communication Technology programme.')
doc.add_paragraph('\n')
add_table('Project approval record', ['Name and role','Signature','Date'], [
    ['Chiranjibi Shrestha\nProject Supervisor',' ',' '],
    ['Binay Malla\nAcademic Coordinator',' ',' '],
    ['Rastra Bhushan Khadka\nPrincipal',' ',' '],
], [2.8,1.8,1.6],10)

# Section 3: remaining front matter.
s3 = doc.add_section(WD_SECTION.NEW_PAGE); configure_section(s3); set_page_numbering(s3,'lowerRoman'); set_footer_page(s3)
add_front_heading('Acknowledgements')
add_body_block('''The completion of SmartBiz was made possible through the guidance, encouragement and practical support of several individuals. I express my sincere gratitude to my project supervisor, Chiranjibi Shrestha, for providing direction throughout the analysis, implementation and documentation stages. His feedback helped transform a broad business-management idea into a focused, testable system.

I am thankful to the faculty members and administration of Virinchi College and Asia e University for providing the academic framework and resources required for this project. I also appreciate the small-business owners, classmates and peers whose observations about daily stock, sales and customer-management difficulties informed the system requirements.

Finally, I acknowledge my family and friends for their patience and continuous encouragement during the twelve-week development period. Their support was especially valuable while integrating multiple services, testing mobile and web workflows, and preparing the final report.''')
add_page_break()
add_front_heading('Abstract')
add_body_block('''Small businesses in Nepal frequently manage products, suppliers, sales, customer credit and leads through notebooks or disconnected spreadsheets. These practices are familiar and inexpensive, but they create duplicate entry, delayed stock updates, limited business visibility and a higher risk of losing important records. SmartBiz was developed as a mobile-first business management system that brings these activities into one authenticated platform while retaining a simple workflow suitable for a small team.

The system uses a Java 21 and Spring Boot 3.4.5 microservice backend, PostgreSQL databases, a Spring Cloud Gateway, Eureka service discovery, Redis caching, a React Native application built with Expo, and a Next.js web dashboard. Its core functions include authentication and billing, inventory and supplier management, point-of-sale recording, sales analytics, customer and lead management, invoice scanning, voice-assisted entry, AI-generated insight cards and controlled sales-file imports. Each service owns its database and communicates with other services through REST APIs. A signed JSON Web Token identifies the user, while the gateway propagates an X-User-Id value so that business records remain isolated by tenant.

The most critical workflow is sale recording. For online payments, SmartBiz reserves stock before checkout, verifies the provider result, and commits the reservation only after successful payment confirmation. Failed, cancelled or expired payments release the reservation and do not deduct stock. This reduces overselling risk without introducing direct cross-service database access. Flyway migrations, DTO-based responses, containerised deployment and lenient Redis error handling support maintainability and recovery.

Repository-level automated tests were executed with the required Java 21 runtime. Fifty-seven of fifty-eight discovered tests passed; one AI insight-card expectation remains a documented quality issue because the generated card set omitted a RESTOCK_SOON signal for the test fixture while returning the other expected business signals. The evaluation therefore concludes that the MVP is functionally comprehensive and architecturally consistent, while production readiness would benefit from resolving that test, formal load testing, refresh-token rotation, stronger observability and expansion of the planned messaging and notification features.''')
add_page_break()
add_front_heading('Table of Contents')
p=doc.add_paragraph(style='SB TOC Field'); add_field(p,'TOC \\o "1-3" \\h \\z \\u','Update table of contents')
add_page_break()
add_front_heading('List of Tables')
p=doc.add_paragraph(style='SB TOC Field'); add_field(p,'TOC \\h \\z \\t "SB Table Caption,1"','Update list of tables')
add_page_break()
add_front_heading('List of Figures')
p=doc.add_paragraph(style='SB TOC Field'); add_field(p,'TOC \\h \\z \\t "SB Figure Caption,1"','Update list of figures')
add_page_break()
add_front_heading('List of Abbreviations')
add_table('Abbreviations used in the report', ['Abbreviation','Meaning'], [
    ['AI','Artificial Intelligence'],['API','Application Programming Interface'],['CRM','Customer Relationship Management'],
    ['DFD','Data-Flow Diagram'],['DTO','Data Transfer Object'],['ERD','Entity Relationship Diagram'],
    ['Expo','Framework and tooling for React Native applications'],['FYP','Final Year Project'],['HTTP/HTTPS','Hypertext Transfer Protocol / Secure HTTP'],
    ['JPA','Java Persistence API'],['JSON','JavaScript Object Notation'],['JWT','JSON Web Token'],
    ['MVP','Minimum Viable Product'],['OAuth','Open Authorization'],['OCR','Optical Character Recognition'],
    ['POS','Point of Sale'],['REST','Representational State Transfer'],['SaaS','Software as a Service'],
    ['SQL','Structured Query Language'],['UAT','User Acceptance Testing'],['UI/UX','User Interface / User Experience'],
    ['UML','Unified Modeling Language'],['URL','Uniform Resource Locator'],
], [1.55,4.7],10)

# Section 4: Chapters 1–3, Arabic numbering.
s4 = doc.add_section(WD_SECTION.NEW_PAGE); configure_section(s4); set_page_numbering(s4,'decimal',1); set_footer_page(s4)

add_heading('CHAPTER 1: INTRODUCTION',1)
add_heading('1.1 Background',2)
add_body_block('''Small enterprises form an important part of Nepal’s local economy. Retail shops, wholesalers, service providers and home-based businesses often operate with a small number of employees and limited administrative resources. The owner commonly performs several roles at once: purchasing goods, serving customers, recording credit, following up with leads and reviewing cash flow. Because daily operations take priority, record keeping is frequently based on paper registers, messaging applications or separate spreadsheets.

These tools can support a business at an early stage, but they do not maintain one dependable view of stock, sales and customer activity. A product may be sold without the notebook being updated, a supplier payment may be recorded in another file, or a promising lead may be forgotten after an informal phone conversation. When the owner later asks a simple question—such as which items should be reordered, which customers have unpaid balances, or whether this week performed better than last week—the answer requires manual calculation.

SmartBiz addresses this problem through a mobile-first system supported by a web dashboard. The mobile application is intended for immediate operational tasks at the counter or storeroom, while the web dashboard supports broader review on a larger screen. Both clients use the same backend services and the same authenticated business data. This shared foundation reduces duplicate entry and permits a sale, stock movement or customer update to become visible across the system.

The project was designed as a twelve-week solo final year project under a fixed technology requirement: Java Spring Boot and PostgreSQL must form the backend foundation. The final implementation extends this requirement with service discovery, an API gateway, React Native, Next.js, Redis, Docker Compose and Gemini-assisted analysis. The architecture is intentionally modular so that critical domains can evolve independently without allowing one service to read another service’s database directly.''')

add_heading('1.2 Problem Statement',2)
add_body('Small businesses need an affordable and understandable way to coordinate products, suppliers, sales, customers and leads. Existing manual or fragmented methods create five connected difficulties:')
add_bullets([
    'Inventory counts can become inaccurate because sales and restocking are recorded at different times or in different places.',
    'Supplier balances, customer due amounts and sales totals require repeated manual reconciliation.',
    'Customer history and lead follow-up information are not consistently available when a decision must be made.',
    'Owners receive little analytical support from raw records and may reorder too late, retain slow-moving stock or miss cash-flow warning signs.',
    'Using separate applications for each activity increases cost, repeated data entry and learning effort for a small team.'
])
add_body('The design problem is therefore not limited to digitising a register. The system must preserve tenant isolation, prevent invalid stock deductions, integrate online payment states, recover when optional infrastructure such as Redis is unavailable, and still present the result through an interface that remains practical on a phone.')

add_heading('1.3 Relevance to Information Technology',2)
add_body_block('''SmartBiz applies several core areas of information technology in one working product. The backend demonstrates distributed service design, service discovery, gateway routing, REST communication, database isolation, transactions, caching and containerisation. The clients demonstrate responsive information architecture, secure token storage, asynchronous API integration and mobile device features such as camera input.

The project also explores the responsible placement of generative AI within a business workflow. Gemini is not treated as the system of record. It receives authorised context through the AI service and produces explanations, parsed proposals or insight cards that can be reviewed before a business record is committed. This distinction between deterministic transaction processing and probabilistic assistance is important for reliable software.

Finally, the project is relevant to cybersecurity and data governance. Authentication is performed at the gateway, user identity is propagated to downstream services, and repositories filter by user ID. API responses use DTOs rather than exposing persistence entities. Environment variables separate credentials from source code, and database changes are versioned through Flyway migrations.''')

add_heading('1.4 Aim and Objectives',2)
add_heading('1.4.1 Main Aim',3)
add_body('The main aim is to design, implement and evaluate a mobile-first business management system that gives a small business one secure platform for daily operations, business records and decision support.')
add_heading('1.4.2 Specific Objectives',3)
add_bullets([
    'Implement secure signup, login, email verification, password recovery, Google OAuth, profile management and tenant-aware access.',
    'Provide product, category, supplier, stock, supplier-ledger and low-stock management.',
    'Implement point-of-sale recording, sales history, weekly trends and online-payment-aware stock reservation.',
    'Provide customer records, purchase totals, due amounts, interactions and a structured lead pipeline with lead-to-customer conversion.',
    'Develop React Native and Next.js clients connected to the same backend through the API gateway.',
    'Use AI for contextual questions, insight cards, invoice interpretation, voice-assisted entry and controlled import reconciliation.',
    'Apply PostgreSQL, Flyway, DTOs, caching, Docker Compose and automated tests to support reliability and maintainability.'
])

add_heading('1.5 Scope of the Project',2)
add_heading('1.5.1 Functional Scope',3)
add_body('The implemented scope covers six business areas and their supporting infrastructure:')
add_table('Table 1.5.1: Functional scope of SmartBiz', ['Area','Included functions'], [
    ['Identity and billing','Signup, login, verification, password reset, Google OAuth, profile, plan status and payment records.'],
    ['Inventory','Categories, products, images, barcode lookup, low stock, adjustments, restock, suppliers, balances, ledger and stock reservations.'],
    ['Sales','POS sale creation, history, import, today/weekly/trend analytics, eSewa workflow and payment-state handling.'],
    ['CRM','Customer CRUD, purchase total, due amount, interactions, lead stages, follow-up and conversion to customer.'],
    ['AI assistance','Business questions, insight cards, invoice scan, voice parsing, sales-file parsing, staged import analysis and reconciliation.'],
    ['User interfaces','Expo mobile application and Next.js web dashboard using shared gateway endpoints.'],
], [1.55,4.7],9.5)
add_heading('1.5.2 Technical Scope',3)
add_body('The technical scope includes Java 21, Spring Boot 3.4.5, Spring Cloud Gateway, Eureka, PostgreSQL, Redis 7, Flyway, Docker Compose, React Native with Expo, TypeScript and Next.js. Managed external services include Neon PostgreSQL, Gemini, Google OAuth and payment-provider interfaces. The repository contains an optional messaging module, but unified inbox behavior is outside the completed core deployment.')

add_heading('1.6 Limitations',2)
add_bullets([
    'The project is an academic MVP developed by one person in twelve weeks; it has not completed a production security audit or a high-volume load test.',
    'Internet connectivity is required for the hosted databases, AI requests, OAuth and online payment verification.',
    'AI responses may be incomplete or inaccurate and therefore remain advisory; users must review proposed imports and business recommendations.',
    'Voice recognition requires a native development build for full device speech support; Expo Go uses a text fallback.',
    'The Stripe path is disabled by default and the eSewa integration uses configurable test/UAT behavior until production credentials and compliance checks are completed.',
    'Unified messaging, Firebase push notifications and refresh-token rotation remain planned enhancements.',
    'The report uses a student-ID placeholder because a verified ID was not available in the project materials.'
])

add_heading('1.7 Target Users',2)
add_body('The primary target user is the owner or manager of a small retail, wholesale or service business in Nepal. Secondary users are trusted staff members who record sales, update stock or follow up with customers. The design assumes limited time for training, frequent phone use and a need for clear summaries rather than complex enterprise configuration.')

add_heading('1.8 Organisation of the Report',2)
add_body('Chapter 1 introduces the problem, objectives and scope. Chapter 2 reviews representative business tools and identifies the gap addressed by SmartBiz. Chapter 3 describes requirements, methodology, feasibility, risks and system design. Chapter 4 explains implementation technologies and module behavior. Chapter 5 presents testing evidence and evaluation. Chapter 6 concludes the project and proposes future improvements, followed by references and appendices.')

add_heading('CHAPTER 2: LITERATURE REVIEW',1)
add_heading('2.1 Introduction',2)
add_body_block('''A literature and product review was undertaken to understand how existing systems address inventory, sales, customer relationships and business analytics. The objective was not to reproduce an enterprise platform, but to identify patterns that are useful for a smaller Nepalese business and to recognise where a focused mobile-first system can reduce complexity.

The review considers four categories: spreadsheet-based record keeping, integrated enterprise suites, specialist inventory or point-of-sale products, and customer relationship platforms. Each category solves part of the problem but introduces trade-offs in cost, configuration, connectivity, learning effort or data fragmentation.''')

add_heading('2.2 Spreadsheet and Paper-Based Management',2)
add_body_block('''Paper registers are inexpensive, familiar and independent of electricity or network availability. They are useful for quick notes, but their information cannot be searched reliably, summarised automatically or shared between devices. Corrections can be difficult to trace and the same figure may be copied into multiple books.

Spreadsheets improve calculation and filtering, yet they still depend on disciplined manual entry. Concurrent edits, inconsistent formulas and uncontrolled copies can create multiple versions of the truth. A spreadsheet does not naturally enforce a stock reservation, authenticate each user, verify payment state or isolate data by business account. SmartBiz therefore retains the simplicity of list-based entry while adding validation, identity and domain rules.''')

add_heading('2.3 Representative Existing Platforms',2)
add_heading('2.3.1 Odoo',3)
add_body('Odoo represents an integrated enterprise suite with modules for inventory, sales, accounting and customer relationships. Its modular scope demonstrates the benefit of shared business data, but configuration and process depth can be disproportionate for a small owner who only needs essential workflows. SmartBiz adopts the integrated-domain idea with a smaller, opinionated feature set.')
add_heading('2.3.2 Zoho Inventory',3)
add_body('Zoho Inventory represents a cloud inventory and order-management product. It illustrates the value of product tracking, supplier information and reports. A specialised inventory product may still require a separate CRM or locally suitable payment workflow, causing the user to move between systems.')
add_heading('2.3.3 Square Point of Sale',3)
add_body('Square POS demonstrates how a mobile-friendly checkout can connect products, payments and receipts. Its strength is transaction speed, but availability, payment support and surrounding business features vary by region. SmartBiz treats POS as one part of a broader inventory, supplier, customer and lead workflow and includes configurable eSewa-oriented behavior for the Nepalese context.')
add_heading('2.3.4 HubSpot CRM',3)
add_body('HubSpot CRM demonstrates structured contacts, leads and pipeline stages. Its approach supports follow-up discipline, but a dedicated CRM does not necessarily maintain the shop’s physical stock or make an inventory reservation when payment begins. SmartBiz connects lead conversion and customer totals to the operational sales domain through service APIs.')

add_heading('2.4 Identified Gaps',2)
add_body('The review identified the following gaps for the target project context:')
add_bullets([
    'A small business may need inventory, POS, customers, leads and simple analytics without the configuration burden of a broad enterprise suite.',
    'International products may not align with Nepalese payment conventions or an owner’s preference for a phone-first workflow.',
    'Specialised products can fragment data and require repeated entry between inventory, CRM and accounting-like records.',
    'Raw dashboards explain what happened but may not translate records into understandable reorder, slow-stock, bundle or cash-flow signals.',
    'Many systems assume a stable connection and mature operational process; an academic MVP must also degrade safely when optional caching or AI services fail.'
])

add_heading('2.5 Proposed System',2)
add_body_block('''SmartBiz combines the minimum set of business domains required for the identified user. Products, suppliers, sales, customers and leads are separate modules in the user interface but share one identity and gateway. The backend separates these domains into services with independently owned data, reducing accidental coupling while keeping integration explicit.

The system adds AI at review points rather than at irreversible transaction boundaries. A scanned invoice or imported file becomes a structured proposal that the user can inspect. Insight cards are derived from business context but do not alter stock. Sales and payments remain deterministic and are protected by reservations and provider verification.

Mobile and web clients are complementary. The mobile application supports quick operational use, camera input and voice assistance. The web dashboard supports paginated tables and wider analytics. This arrangement responds to the target user’s need for mobility without preventing office-based review.''')

add_table('Table 2.5.1: Comparative view of representative approaches', ['Criterion','Paper / spreadsheet','Specialist product','Enterprise suite','SmartBiz'], [
    ['Integrated inventory, sales and CRM','Low','Usually partial','High','High within MVP scope'],
    ['Mobile-first daily workflow','Manual / limited','Often high','Varies','Core design goal'],
    ['Nepal-oriented payment path','Manual','Region dependent','Requires configuration','Configurable eSewa flow'],
    ['Tenant-aware API architecture','Not applicable','Vendor managed','Vendor managed','JWT + X-User-Id'],
    ['AI-assisted scan/import','No','Varies','Varies','Review-before-commit workflow'],
    ['Operational complexity','Low initially; grows with data','Medium','High','Moderate and focused'],
    ['Offline capability','Paper: high','Varies','Varies','Limited in current MVP'],
], [1.25,1.25,1.25,1.25,1.25],8.7)

add_heading('2.6 Conceptual Foundations',2)
add_heading('2.6.1 Microservice Boundaries',3)
add_body('A service boundary should correspond to a cohesive business responsibility and should own its persistence model. SmartBiz applies this principle to identity/billing, inventory, CRM, sales and AI. The gateway and discovery server support communication but do not own business records. Direct cross-service database access is prohibited; shared identifiers are exchanged through REST.')
add_heading('2.6.2 Transactional Integrity',3)
add_body('PostgreSQL transactions provide all-or-nothing behavior inside one service database. SmartBiz uses this property for stock updates and reservation commits. Since a database transaction cannot safely span independent services without a distributed transaction mechanism, the sale workflow uses explicit states: reserve, verify, commit or release. This is a compensating workflow rather than a hidden cross-database transaction.')
add_heading('2.6.3 Cache-Aside Resilience',3)
add_body('Redis reduces repeated work for paginated list endpoints, but it is treated as an optimisation. Cache configuration uses a lenient error handler so a stale or malformed entry produces a warning and the service falls through to PostgreSQL. Write operations evict affected collections. This preserves correctness when Redis is temporarily unavailable.')
add_heading('2.6.4 Human-in-the-Loop AI',3)
add_body('Generative AI is useful for interpreting unstructured input and explaining patterns, but its output is probabilistic. SmartBiz therefore separates an AI proposal from an authoritative business transaction. Import sessions record state, analysis and reconciliation before commit, providing a review boundary that supports accountability.')

add_heading('2.7 Chapter Summary',2)
add_body('The review shows that existing approaches offer valuable individual capabilities but can be too fragmented or too complex for the target user. SmartBiz’s contribution is a focused combination of mobile operations, shared data, explicit service boundaries, safe payment-linked stock handling and reviewable AI assistance.')

add_heading('CHAPTER 3: SYSTEM ANALYSIS AND DESIGN',1)
add_heading('3.1 Introduction',2)
add_body('This chapter translates the problem into requirements and describes the design selected to satisfy them. It covers stakeholders, feasibility, methodology, risks, architecture, data ownership, data flow, use cases, the critical sale sequence and interface planning.')

add_heading('3.2 Stakeholder and Requirement Analysis',2)
add_heading('3.2.1 Stakeholders',3)
add_table('Table 3.2.1: Stakeholders and interests', ['Stakeholder','Primary interest'], [
    ['Business owner','Accurate stock, sales visibility, customer credit, low learning effort and control over access.'],
    ['Staff user','Fast entry, clear validation and predictable task flow on mobile.'],
    ['Customer','Accurate sale, payment status and due records.'],
    ['Supplier','Consistent purchase, payment and adjustment ledger information.'],
    ['Project evaluator','Evidence that requirements, design, implementation and testing are aligned.'],
    ['System administrator / developer','Maintainable services, migrations, logs, configuration and recovery behavior.'],
], [1.9,4.35],9.5)
add_heading('3.2.2 Functional Requirements',3)
add_table('Table 3.2.2: Prioritised functional requirements', ['ID','Requirement','Priority'], [
    ['FR-01','Register, verify and authenticate a user; recover access and update profile.','Must'],
    ['FR-02','Create, read, update, search and delete products and categories.','Must'],
    ['FR-03','Adjust stock, show low-stock products and maintain supplier balances/ledger.','Must'],
    ['FR-04','Create a sale with validated quantities and maintain sale history.','Must'],
    ['FR-05','Reserve stock during online checkout and commit only after verified success.','Must'],
    ['FR-06','Manage customers, due amounts, interactions, leads and lead conversion.','Must'],
    ['FR-07','Provide weekly/today/trend analytics and dashboard summaries.','Should'],
    ['FR-08','Answer contextual questions and produce reviewable AI insight cards.','Should'],
    ['FR-09','Parse invoice images, voice text and sales files into proposals.','Should'],
    ['FR-10','Support plan/billing state and provider callbacks.','Could'],
], [.7,4.7,.85],9.2)
add_heading('3.2.3 Non-Functional Requirements',3)
add_table('Table 3.2.3: Non-functional requirements', ['Category','Requirement and design response'], [
    ['Security','Signed JWT, gateway validation, tenant header, user-filtered repository queries, validation, DTO responses and externalised secrets.'],
    ['Reliability','Atomic local transactions, reservation state machine, idempotent payment checks, cache fall-through and explicit error responses.'],
    ['Performance','Pagination, Redis caching, compact DTOs and independent scaling of service containers.'],
    ['Usability','Mobile-first navigation, search/filter patterns, status badges, review screens and consistent feedback.'],
    ['Maintainability','Layered services, domain DTOs, Flyway migrations, TypeScript service modules, shared tokens and Docker configuration.'],
    ['Portability','Java 21 containers, environment variables and Docker Compose across supported hosts.'],
    ['Auditability','Persisted sale/payment/import states and append-like supplier/customer interaction records.'],
], [1.25,5.0],9.2)

add_heading('3.3 Development Methodology',2)
add_body_block('''An incremental methodology was selected because the project combined several dependent domains and had a fixed twelve-week schedule. Work was divided into short vertical slices: infrastructure and identity, inventory, CRM, sales, clients, AI assistance, billing, caching and evaluation. A slice was considered usable only when its endpoint, persistence, client integration and error handling could be exercised together.

This approach reduced the risk of producing isolated screens or endpoints. For example, the inventory feature was not treated as complete when product CRUD existed; it was extended through mobile/web integration, pagination, supplier behavior and stock reservation. Git provided version history, while Docker Compose provided a repeatable integration environment for all services.''')
add_table('Table 3.3.1: Incremental project phases', ['Phase','Major outputs'], [
    ['Foundation','Requirements, architecture, repository structure, Eureka, gateway and environment strategy.'],
    ['Core domains','Authentication, inventory/suppliers, CRM/leads and database migrations.'],
    ['Transactions','Sales, stock checks, reservation/commit/release and analytics.'],
    ['Experience','Expo mobile screens, Next.js dashboard, pagination and reusable UI patterns.'],
    ['Intelligence and billing','Gemini workflows, import sessions, insight cards, OAuth and payment providers.'],
    ['Quality and delivery','Caching, exception handling, tests, Docker builds, documentation and evaluation.'],
], [1.7,4.55],9.5)

add_heading('3.4 Feasibility Study',2)
add_heading('3.4.1 Technical Feasibility',3)
add_body('The chosen technologies are compatible with the project constraint and available development environment. Java 21 supports the Spring Boot services, PostgreSQL is available locally or through Neon, and Expo supports development on Windows with Android devices. Docker Compose can start the distributed backend from one configuration. The principal technical risks were service integration, payment state and native speech support; each received a fallback or explicit state model.')
add_heading('3.4.2 Operational Feasibility',3)
add_body('The workflows mirror familiar small-business activities: add product, record sale, update customer and review totals. Search, filters, status badges and quick actions reduce the number of steps. The mobile application does not require a user to understand the underlying microservices. Operational adoption still depends on initial product entry and consistent use at the point of sale.')
add_heading('3.4.3 Economic Feasibility',3)
add_body('The project uses open-source frameworks and can run in low-cost containers. Managed PostgreSQL, AI and payment services may introduce usage-based costs, but the architecture permits local development and configurable integrations. For an academic MVP, the primary cost was development time rather than software licensing.')
add_heading('3.4.4 Schedule Feasibility',3)
add_body('The scope was organised around a twelve-week plan. Core inventory, sales, CRM and authentication were prioritised before optional messaging and push notifications. This prioritisation enabled a complete business workflow while moving lower-priority features to the future roadmap.')

add_heading('3.5 Data Collection and Analysis',2)
add_body('Requirements were derived from observation of common retail record-keeping tasks, informal discussions about stock and credit tracking, inspection of representative software categories and continuous review of the working prototype. The analysis focused on repeated pain points rather than attempting a statistically generalisable market study. Therefore, conclusions are appropriate for project design but should be validated with structured field research before commercial deployment.')

add_heading('3.6 SWOT and Risk Analysis',2)
add_table('Table 3.6.1: SWOT analysis', ['Strengths','Weaknesses','Opportunities','Threats'], [
    ['Integrated mobile/web workflow; modular services; tenant isolation; reviewable AI; payment-aware stock reservation.','Internet dependency; solo-project test capacity; operational complexity of microservices; limited offline support.','Nepal-focused payments; notifications; messaging; richer forecasting; multi-branch support; SaaS deployment.','Credential leakage; provider changes; AI rate limits; user resistance; data quality; network outages.'],
], [1.56,1.56,1.56,1.56],8.7)
add_table('Table 3.6.2: Principal risks and mitigations', ['Risk','Likelihood / impact','Mitigation'], [
    ['Cross-user data exposure','Low / Critical','Gateway identity propagation plus user-scoped repository queries and endpoint tests.'],
    ['Stock deducted for failed payment','Medium / High','Reserve first; verify provider; commit on success; release on failure/expiry.'],
    ['Redis outage or malformed entry','Medium / Medium','Lenient cache error handler and database fall-through.'],
    ['AI hallucination or parse error','Medium / High','Review-before-commit import sessions and deterministic validation.'],
    ['External API rate limit','Medium / Medium','Clear errors, configurable Gemini model/key and non-AI core workflows.'],
    ['Schedule expansion','High / Medium','Prioritised MVP; defer messaging, push and refresh-token work.'],
], [1.85,1.45,2.95],9.2)

add_heading('3.7 System Architecture',2)
add_body('Both user interfaces send HTTPS/JSON requests to the API gateway. The gateway validates JWTs, derives the authenticated user and routes requests to services discovered through Eureka. Each domain service owns a PostgreSQL database. Redis caches selected paginated reads, and the AI service is the only component that calls Gemini. This arrangement is shown in Figure 3.7.1.')
add_figure('smartbiz-system-architecture.png','Figure 3.7.1: SmartBiz system architecture',6.25)
add_body('The architecture deliberately separates availability concerns. A cache failure should not become a data failure, and an AI failure should not prevent product or sale management. The gateway provides one external entry point while the internal network retains named service endpoints. Shared payment DTOs and signatures are placed in a small common library without sharing domain persistence.')

add_heading('3.8 Data Design',2)
add_body('The database-per-service strategy means there is no single global schema. Foreign keys are used only within the database that owns both entities. A sales record can carry a customer ID or product ID received through an API, but the sales database cannot enforce a foreign key into the CRM or inventory database. Service logic validates and reconciles these references.')
add_figure('smartbiz-distributed-data-model.png','Figure 3.8.1: Distributed data model and service ownership',6.25)
add_table('Table 3.8.1: Data ownership by service', ['Service','Owned records','Selected integrity rules'], [
    ['Auth/Billing','Users, verification/reset tokens, plans, subscriptions and provider payments.','Unique identity, token expiry and controlled subscription state.'],
    ['Inventory','Categories, products, suppliers, ledger, images and stock reservations.','Non-negative quantities, user ownership and reservation status.'],
    ['CRM','Customers, interactions and leads.','Tenant filtering, lead-stage values and conversion rules.'],
    ['Sales','Sales, sale items, analytics inputs and payment intents.','Calculated totals, sale status and provider verification state.'],
    ['AI','Import sessions, artifacts and reconciliation metadata.','State transitions and user-approved commit boundary.'],
], [1.15,2.55,2.55],9.1)

add_heading('3.9 Data-Flow Design',2)
add_heading('3.9.1 Context Diagram',3)
add_body('At context level, the business owner interacts with SmartBiz as one system. Customers provide sale/payment information, while external providers supply identity, payment and AI responses.')
add_figure('smartbiz-dfd-level-0.png','Figure 3.9.1: SmartBiz context diagram (DFD Level 0)',6.1)
add_heading('3.9.2 Level-1 Data Flow',3)
add_body('The level-1 view separates identity, inventory, CRM, sales, AI and import processes. Data stores are independently owned and external providers are reached through controlled service calls.')
add_figure('smartbiz-dfd-level-1.png','Figure 3.9.2: SmartBiz data-flow diagram (DFD Level 1)',6.15)

add_heading('3.10 Use-Case Design',2)
add_body('The business owner can access all operational and administrative features. A staff user performs routine stock, sales and customer activities subject to the account’s authorization rules. Payment providers and Gemini participate only in bounded use cases.')
add_figure('smartbiz-use-case-overview.png','Figure 3.10.1: SmartBiz use-case overview',6.1)

add_heading('3.11 Critical Sale and Payment Sequence',2)
add_body_block('''A direct cash sale can validate stock and commit the sale without an external checkout. An online payment requires a longer sequence. The sales service first requests an inventory reservation. A provider intent is then created and the customer completes the external flow. The callback alone is not trusted; the sales service verifies provider status before committing the reservation.

If the provider reports failure, cancellation or expiry, the reservation is released. The stock deduction is atomic inside the inventory database, but the complete cross-service workflow is not described as one distributed database transaction. Its safety comes from explicit state transitions, verification and compensation.''')
add_figure('smartbiz-sale-payment-sequence.png','Figure 3.11.1: Record-sale and online-payment sequence',6.25)

add_heading('3.12 Deployment Design',2)
add_body('The backend is packaged as containers on an internal Compose network. Mobile and web clients access the public gateway, while PostgreSQL, Gemini and payment providers remain external managed services. Health checks and environment variables support repeatable startup and configuration.')
add_figure('smartbiz-deployment-topology.png','Figure 3.12.1: SmartBiz deployment topology',6.2)

add_heading('3.13 Interface Design',2)
add_body_block('''The interface uses repeated patterns to reduce learning effort: search at the top, filters near the result list, status badges for state, expandable detail cards on mobile and paginated tables on web. Primary actions are visually prominent, while destructive actions require deliberate selection.

The mobile navigation exposes the most frequently used business domains as tabs. Camera and voice actions appear in context instead of as separate applications. The web dashboard preserves the same concepts but uses wider tables, previous/next pagination and larger analytics surfaces.''')
add_figure('smartbiz-interface-wireframes.png','Figure 3.13.1: Representative mobile and web wireframes',6.25)
add_heading('3.14 Chapter Summary',2)
add_body('The design connects usability requirements to explicit technical controls. The next chapter explains how these designs were implemented in the repository.')

# Section 5: Chapters 4–5.
s5 = doc.add_section(WD_SECTION.NEW_PAGE); configure_section(s5); set_page_numbering(s5,'decimal'); set_footer_page(s5)

add_heading('CHAPTER 4: IMPLEMENTATION',1)
add_heading('4.1 Introduction',2)
add_body('Implementation followed the service boundaries and incremental plan established in Chapter 3. This chapter describes technologies, repository organisation, backend modules, client applications, security, caching, AI integration, deployment and project resources.')

add_heading('4.2 Technologies and Tools',2)
add_table('Table 4.2.1: Main implementation technologies', ['Layer','Technology','Purpose in SmartBiz'], [
    ['Backend language/runtime','Java 21','Required runtime for all Spring services and shared payment library.'],
    ['Backend framework','Spring Boot 3.4.5','REST controllers, validation, transactions, security, data access and configuration.'],
    ['Gateway/discovery','Spring Cloud Gateway and Eureka','Single entry point, routing, JWT identity propagation and service lookup.'],
    ['Database','PostgreSQL 15+ / Neon','Relational persistence with transactions and per-service databases.'],
    ['Migration','Flyway','Versioned schema creation and forward-only changes.'],
    ['Cache','Redis 7','Cached paginated reads with database fall-through on errors.'],
    ['Mobile','React Native 0.81.5, Expo 54, TypeScript','Phone-first operational interface and device features.'],
    ['Web','Next.js 16.2.4, React 19.2.4, Tailwind CSS 4','Server/client dashboard, tables and analytics interface.'],
    ['AI','Gemini API','Contextual chat, insight cards, scan/voice/file interpretation.'],
    ['Deployment','Docker and Docker Compose','Multi-stage images, internal networking and repeatable startup.'],
], [1.25,2.05,2.95],8.8)

add_heading('4.3 Repository Organisation',2)
add_code_block('''smartBiz/
  backend/
    eureka-server/        api-gateway/
    auth-service/         inventory-service/
    crm-service/          sales-service/
    ai-service/           payment-common/
  mobile/                 Expo / React Native application
  frontend/web/           Next.js dashboard
  docker-compose.yml      project documentation''')
add_body('Within each Spring service, controllers expose DTO-based APIs, services implement business rules, repositories isolate persistence, configuration classes define security/caching behavior and Flyway scripts version the schema. Mobile and web clients each use a service layer so UI components do not construct authentication headers independently.')

add_heading('4.4 Backend Infrastructure',2)
add_heading('4.4.1 Eureka Server',3)
add_body('Eureka runs on port 8761 and provides registration and lookup for named services. This allows gateway routes to target logical service names instead of fixed container IP addresses. Health and registration behavior are part of the Compose startup sequence.')
add_heading('4.4.2 API Gateway',3)
add_body('The gateway runs on port 8080 and routes /auth and /billing to Auth, /inventory to Inventory, /customers and /leads to CRM, /sales to Sales, and /ai to the AI service. Its authentication filter validates the bearer token, extracts the user identity and supplies X-User-Id to downstream services. Public authentication and provider-callback routes are handled explicitly.')

add_heading('4.5 Auth and Billing Service',2)
add_body_block('''The Auth service stores users and account-related state. Registration validates the submitted identity and password, while login returns the token used by both clients. Email verification, resend, forgot-password and reset-password flows use time-limited tokens. Google OAuth provides an alternative identity path. Profile updates return DTOs rather than the user entity.

Billing endpoints expose plan, status, checkout and payment history behavior. Provider-specific details are kept behind service interfaces. eSewa and Stripe configurations are environment-driven and may remain disabled or in test mode when credentials are not supplied. Callback handlers validate provider information before changing payment or subscription state.''')

add_heading('4.6 Inventory and Supplier Service',2)
add_body_block('''Inventory provides paginated product and supplier endpoints as well as category management. Product records include user ownership, category, stock, price, barcode and optional image information. Search and filters are applied within the user’s data set. Signed image operations prevent an upload from becoming a permanent product attachment until the update succeeds.

Supplier management includes summary values, payment and adjustment operations and a ledger. This makes supplier balance changes traceable instead of storing only an unexplained total. Stock endpoints support manual adjustment, restock and low-stock queries.

Reservations are central to safe online sales. A reservation temporarily holds the requested quantity against a sale reference and has an expiry/status. Commit performs an atomic inventory update; release returns the held quantity to availability. Internal endpoints require trusted service context rather than being exposed as normal user operations.''')

add_heading('4.7 CRM Service',2)
add_body_block('''CRM owns customers, interactions and leads. Customer list endpoints are paginated and tenant-filtered. Purchase totals and due amounts can be updated through controlled endpoints used by the sales workflow. Interactions record the type and notes for later review.

Leads progress through NEW, CONTACTED, INTERESTED, PROPOSAL, WON or LOST. Source, estimated value, notes and follow-up date support basic pipeline management. Conversion creates a customer from a qualified lead and removes the lead inside one CRM transaction. The conversion does not require another service database.''')

add_heading('4.8 Sales Service',2)
add_body_block('''Sales receives a user-scoped sale request, validates its items and coordinates required inventory and CRM calls. It stores the sale and item values needed for history and analytics. Today, weekly and trend endpoints aggregate authoritative sale data rather than relying on client-side totals.

For online eSewa flow, the service creates an intent only after inventory can be reserved. Provider callbacks are followed by status verification. A verified success commits stock and confirms the sale; failure, cancellation or expiry releases the reservation. Retry and state checks reduce duplicate effects. Customer totals or due amounts are updated through CRM APIs after the sale state is known.

File-imported sales enter through validated import endpoints. The client or AI service may propose rows, but the sales service remains responsible for final domain validation and persistence.''')

add_heading('4.9 AI Service',2)
add_body_block('''The AI service obtains authorised inventory, sales and CRM context through service APIs and constructs a bounded prompt for Gemini. It supports general business questions and structured insight cards such as slow-moving stock, bundle opportunities, cash-flow warnings and restock signals. The service does not query another service’s database.

Invoice and voice endpoints transform unstructured input into proposed products, leads or chat text. Sales-file import uses sessions, artifacts, analysis, reconciliation, commit and close steps. This staged design gives the user an opportunity to review ambiguous mappings before records are written. Gemini 2.5 Flash Lite is configurable because earlier model availability and regional quota behavior can vary.''')

add_heading('4.10 Mobile Application',2)
add_body('The Expo application uses TypeScript and file-based routing. SecureStore retains authentication data, and a shared Axios interceptor attaches Authorization and X-User-Id values. The principal tabs cover home, inventory, suppliers, sales, customers, leads, AI and settings. Product entry, POS cart behavior, lead stages, customer history, analytics and account changes all call the gateway. Camera, image picker and speech-recognition integrations are wrapped with review or fallback behavior.')
add_table('Table 4.10.1: Mobile navigation and responsibilities', ['Tab','Principal responsibilities'], [
    ['Home','Weekly revenue, alerts, quick actions and on-demand AI insight.'],
    ['Inventory','Products, search, stock actions, invoice scan and voice-assisted product entry.'],
    ['Suppliers','Supplier details, balance status, ledger and adjustments.'],
    ['Sales','POS cart, customer/payment selection, receipt state and history.'],
    ['Customers','Customer cards, due values and interaction/history review.'],
    ['Leads','Stage filters, details, follow-up and conversion.'],
    ['AI','Contextual chat plus camera, voice and import entry points.'],
    ['Settings','Profile and account preferences.'],
], [1.35,4.9],9.3)

add_heading('4.11 Web Dashboard',2)
add_body('The Next.js dashboard implements the same major domains for a larger screen. Server pages read pagination values from asynchronous search parameters and request fifteen-row pages. A shared Pagination component renders Previous and Next links and hides when only one page exists. Inventory, suppliers, customers and leads use paginated tables; sales, analytics, AI and billing pages use the same authenticated service layer.')

add_heading('4.12 Security and Tenant Isolation',2)
add_body('Security is enforced as a chain rather than a single check:')
add_bullets([
    'The Auth service issues signed, expiring JWTs after successful identity validation.',
    'The gateway rejects invalid protected requests and derives X-User-Id from the verified token.',
    'Each service receives user identity and repository methods include the user ID in lookups and list queries.',
    'DTOs restrict response fields and request validation rejects structurally invalid input.',
    'Internal service endpoints use trusted tokens/headers and are not presented as public user APIs.',
    'Secrets, database URLs and provider credentials are supplied through environment variables and are not printed in this report.',
    'Provider callbacks and payment status are verified before stock or subscription effects are committed.'
])
add_body('A remaining production improvement is refresh-token rotation with revocation and tighter issuer/audience enforcement across all services. HTTPS termination, secure cookie decisions for web OAuth, rate limiting and an external security review are also part of the hardening roadmap.')

add_heading('4.13 Pagination and Redis Caching',2)
add_body_block('''Large list endpoints return a PagedResponse containing content, current page, total pages, total elements and hasNext. This contract is used by both mobile load-more behavior and web previous/next navigation. Page and size are included in cache keys so one tenant or page cannot return another tenant’s result.

Write methods evict affected caches. Both inventory and CRM cache configurations implement a lenient error handler: read/write/evict failures are logged as warnings and the request continues to the database. Redis therefore improves response efficiency without becoming a required source of truth.''')

add_heading('4.14 API Design and Error Handling',2)
add_body('Endpoints use resource-oriented paths and JSON DTOs. Common operations map to GET, POST, PUT and DELETE, while domain actions such as /convert, /reserve, /commit or /release use explicit POST endpoints. Global exception handlers translate not-found, validation, conflict and internal failures into appropriate status codes and concise response bodies. Clients display recoverable messages rather than exposing stack traces.')
add_table('Table 4.14.1: Representative API routes', ['Method and route','Service','Purpose'], [
    ['POST /auth/login','Auth','Authenticate and issue access token.'],
    ['GET /inventory/products?page=&size=','Inventory','Return tenant-scoped product page.'],
    ['POST /inventory/reservations','Inventory','Reserve stock for a sale reference.'],
    ['POST /leads/{id}/convert','CRM','Create a customer from a qualified lead.'],
    ['POST /sales','Sales','Validate and record a sale.'],
    ['POST /sales/payments/esewa','Sales','Create payment intent after reservation.'],
    ['GET /sales/analytics/weekly','Sales','Return weekly aggregate values.'],
    ['POST /ai/import-sessions','AI','Begin reviewable sales import.'],
], [2.8,1.15,2.3],9.1)

add_heading('4.15 Deployment and Configuration',2)
add_body('Docker Compose builds and starts Eureka, the gateway, core services and Redis on a shared network. Multi-stage images reduce runtime contents. Environment variables provide database URLs, credentials, JWT secret, Eureka URL, Gemini key, OAuth values and provider configuration. Spring configuration includes safe local fallbacks where appropriate, while production secrets must be supplied by the deployment platform.')

add_heading('4.16 Schedule and Resource Use',2)
add_figure('smartbiz-development-gantt.png','Figure 4.16.1: Twelve-week SmartBiz development plan',6.25)
add_table('Table 4.16.1: Project resources', ['Resource type','Items used'], [
    ['Hardware','Windows development computer, Android test device and internet connection.'],
    ['Development software','JDK 21, Maven, Node.js/npm, code editor, Git, Docker Desktop and PostgreSQL tooling.'],
    ['Frameworks','Spring Boot/Cloud, React Native/Expo, Next.js, Redis and Flyway.'],
    ['Hosted services','Neon PostgreSQL, Gemini API, Google OAuth and payment-provider test environments.'],
    ['Documentation','Project context, change log, source comments, framework documentation and this report.'],
], [1.55,4.7],9.4)

add_heading('4.17 Chapter Summary',2)
add_body('The implementation realises the planned domain boundaries and connects mobile and web clients through one authenticated gateway. It also introduces resilience patterns—pagination, cache fall-through, reservations and reviewable AI—that are directly tied to the requirements.')

add_heading('CHAPTER 5: TESTING AND EVALUATION',1)
add_heading('5.1 Introduction',2)
add_body('Testing evaluated individual service rules, client-facing workflows, cross-service integration assumptions and deployment behavior. This chapter reports observed evidence without inventing performance figures that were not measured.')

add_heading('5.2 Testing Strategy',2)
add_body('The strategy combined the following levels:')
add_bullets([
    'Unit tests for payment signatures, controllers and service business rules.',
    'Repository/service tests for user isolation, pagination, reservations and conversion behavior.',
    'Integration-oriented tests using mocked service/provider boundaries for sale and payment state.',
    'Manual workflow checks through the mobile and web clients against gateway endpoints.',
    'Container startup and service-registration checks through Docker Compose.',
    'Documented negative cases for invalid input, insufficient stock, failed payment, cache failure and AI/provider failure.'
])

add_heading('5.3 Automated Test Execution',2)
add_body_block('''Tests were executed with the project’s required Java 21 runtime. The repository contained 58 discovered @Test methods across the shared payment library and five core business services. Fifty-seven passed and one failed. An initial attempt under Java 25 produced a Mockito/Byte Buddy compatibility failure; this was an environment mismatch rather than an application result, and the suite was rerun correctly under Java 21.

The remaining failure occurs in the AI insight-card test. For its fixture, the service returned SLOW_MOVING_STOCK, BUNDLE_OPPORTUNITY and CASH_FLOW_WARNING, while the test also expected RESTOCK_SOON. This is recorded as an unresolved inconsistency between the fixture/expectation and current restock-selection logic. The failed assertion should be resolved before claiming a completely green build.''')
add_table('Table 5.3.1: Automated test result by module', ['Module','Discovered','Passed','Failed','Observation'], [
    ['payment-common','1','1','0','Payment signature behavior passed.'],
    ['auth-service','12','12','0','Authentication, billing and controller rules passed.'],
    ['inventory-service','21','21','0','Product, supplier, stock and reservation cases passed.'],
    ['crm-service','2','2','0','Customer/lead service cases passed.'],
    ['sales-service','19','19','0','Sale, payment and integration-oriented rules passed.'],
    ['ai-service','3','2','1','Restock insight-card expectation remains unresolved.'],
    ['Total','58','57','1','98.3% of discovered tests passed.'],
], [1.35,.75,.75,.65,2.75],8.8)

add_heading('5.4 Functional Test Cases',2)
add_table('Table 5.4.1: Representative functional tests', ['ID','Scenario','Expected result','Result'], [
    ['TC-01','Valid login','JWT and user profile are returned.','Pass'],
    ['TC-02','Request another user’s product','No record is returned or access is rejected.','Pass'],
    ['TC-03','Create product and restock','Product persists and stock increases by validated amount.','Pass'],
    ['TC-04','Sale with insufficient stock','Sale is rejected; stock remains unchanged.','Pass'],
    ['TC-05','Successful online payment','Reservation is verified and committed; sale confirms.','Pass'],
    ['TC-06','Cancelled/failed payment','Reservation is released; stock is not deducted.','Pass'],
    ['TC-07','Convert lead to customer','Customer is created and lead removed in CRM transaction.','Pass'],
    ['TC-08','Malformed Redis entry','Warning is logged and database response is used.','Pass by design/test coverage'],
    ['TC-09','AI import proposal','Rows remain reviewable until explicit commit.','Pass'],
    ['TC-10','Insight-card restock fixture','Expected card set includes RESTOCK_SOON.','Fail—known issue'],
], [.65,1.85,2.75,1.0],8.7)

add_heading('5.5 Security and Isolation Evaluation',2)
add_body_block('''The strongest security property demonstrated by the design is consistent tenant scoping. The gateway derives user identity from the token rather than accepting an arbitrary client-supplied tenant as authoritative. Downstream queries include the user ID, and cross-service calls propagate trusted identity context.

The assessment also identified areas for production hardening. Shared-secret JWT validation requires careful key management; issuer, audience and expiration claims should be enforced consistently. Refresh-token rotation and revocation are not implemented. Rate limiting, structured security logs, dependency scanning, external penetration testing and production TLS/header configuration should be completed before real financial or personally identifiable data is onboarded.''')

add_heading('5.6 Reliability and Failure Evaluation',2)
add_table('Table 5.6.1: Failure behavior evaluation', ['Failure','Designed behavior','Evaluation'], [
    ['Insufficient stock','Reservation or sale request fails before commit.','Prevents negative stock in normal workflow.'],
    ['Provider cancellation/failure','Release reservation and keep sale unconfirmed.','Separates checkout attempt from inventory effect.'],
    ['Duplicate callback/retry','Check persisted state before applying effect again.','Reduces duplicate commit risk.'],
    ['Redis unavailable/corrupt','Log warning and query PostgreSQL.','Cache does not become a single point of failure.'],
    ['Gemini rate limit/error','Return controlled error; core CRUD/POS remains usable.','AI is optional to system-of-record tasks.'],
    ['Downstream service unavailable','Return explicit failure; do not silently fabricate success.','Requires retry/observability improvements for production.'],
], [1.55,2.55,2.15],9.1)

add_heading('5.7 Performance Evaluation',2)
add_body_block('''Formal concurrent load testing was not completed, so the report does not claim a measured response-time or throughput target. Performance readiness was instead evaluated through structural controls: paginated endpoints bound response size, Redis avoids repeated list work, database indexes accompany migration design, and service containers can be monitored or scaled separately.

The next evaluation stage should establish a reproducible data set and run gateway-level tests for p50, p95 and p99 latency, error rate, throughput and resource use. Tests should cover cold/warm cache behavior, a Redis outage, concurrent reservations for the final stock unit, payment callback bursts and AI requests under provider limits.''')

add_heading('5.8 Objective Achievement',2)
add_table('Table 5.8.1: Evaluation against objectives', ['Objective','Evidence','Status'], [
    ['Secure tenant-aware identity','Auth flows, gateway validation and user-scoped repositories.','Achieved for MVP'],
    ['Inventory and supplier control','CRUD, search, low stock, ledger, images and reservations.','Achieved'],
    ['Sales and analytics','POS, payment states, history and aggregate endpoints.','Achieved'],
    ['CRM and lead pipeline','Customers, due, interactions, stages and conversion.','Achieved'],
    ['Mobile and web access','Expo mobile tabs and Next.js dashboard use shared APIs.','Achieved'],
    ['AI-assisted operations','Chat, cards, invoice/voice/file parsing and staged import.','Achieved with known test issue'],
    ['Maintainable deployment','Flyway, DTOs, caching, Docker Compose and tests.','Achieved for academic MVP'],
], [1.7,3.45,1.1],8.9)

add_heading('5.9 Challenges Encountered',2)
add_bullets([
    'Coordinating stock, sale, customer and payment state without sharing databases required an explicit reservation workflow.',
    'Gateway routes and service security matchers had to remain aligned when adding new resource paths.',
    'Spring Boot does not automatically load a root .env file during local Maven execution, while Docker Compose does; configuration needed clear instructions and fallbacks.',
    'Gemini regional model availability and rate limits required a configurable model and graceful failure handling.',
    'Native speech recognition is unavailable in Expo Go, requiring a development build and a text fallback.',
    'Stale or malformed cache entries initially risked repeated failures; lenient cache error handlers were introduced.',
    'Test execution under the wrong Java version produced tooling incompatibility; the authoritative run used Java 21.'
])

add_heading('5.10 Discussion',2)
add_body_block('''The evaluation indicates that SmartBiz meets its main academic objective: it provides one coherent workflow across the principal records of a small business. The strongest aspect is not the number of screens, but the alignment between UI actions and backend controls. A payment attempt does not automatically become a stock deduction, a cache entry is not treated as authoritative data, and an AI interpretation is not committed without review.

The microservice design demonstrates relevant distributed-system concepts, although it carries more operational complexity than a monolith would for a single small business. For a commercial product, the separation would be justified by SaaS growth, independent scaling and team ownership. For the academic context, it provides a valuable implementation of boundaries and inter-service consistency, but deployment and observability require further maturity.

The known AI test failure prevents a claim of perfect automated verification. Its documentation is itself important: quality reporting should reflect actual evidence. The issue is bounded to the composition of advisory insight cards and does not indicate a failed stock or payment transaction. Nevertheless, it should be investigated and the full build rerun before release.''')

add_heading('5.11 Chapter Summary',2)
add_body('Testing confirms broad functional coverage and validates the most critical inventory, sales, authentication and CRM rules. The outstanding AI assertion, absent load benchmark and production-hardening items define a clear next quality stage.')

# Section 6: conclusion, references and appendices.
s6 = doc.add_section(WD_SECTION.NEW_PAGE); configure_section(s6); set_page_numbering(s6,'decimal'); set_footer_page(s6)

add_heading('CHAPTER 6: CONCLUSION AND FUTURE ENHANCEMENTS',1)
add_heading('6.1 Conclusion',2)
add_body_block('''SmartBiz was developed to replace fragmented small-business records with one mobile-first system supported by a web dashboard. The completed MVP integrates identity, billing, inventory, suppliers, sales, payments, analytics, customers, leads and AI-assisted workflows through a Spring Boot and PostgreSQL architecture.

The project demonstrates that reliable business behavior depends on the relationships between modules. Tenant identity is carried from the gateway into every query. Stock is reserved before online checkout and committed only after verification. Customer totals are updated through CRM APIs rather than shared tables. Redis improves list performance but can fail without blocking database access. Gemini interprets and explains information but is separated from authoritative transaction rules.

The mobile application makes routine operations available at the point of work, while the web dashboard supports wider review. Flyway migrations, DTOs, containers, pagination and automated tests improve maintainability. The observed result—57 of 58 tests passing under Java 21—shows strong coverage for an academic MVP while transparently identifying one remaining AI insight-card inconsistency.

Overall, SmartBiz achieves its central aim and provides a credible foundation for further research and product development. It is not yet a production financial platform, but its architecture and completed workflows demonstrate how a focused system can make small-business information more consistent, visible and actionable.''')

add_heading('6.2 Future Enhancements',2)
add_heading('6.2.1 Immediate Quality Improvements',3)
add_bullets([
    'Resolve the RESTOCK_SOON insight-card fixture/logic mismatch and restore a fully green test run.',
    'Add gateway-level integration tests and concurrent reservation tests with real PostgreSQL and Redis containers.',
    'Introduce structured logs, correlation IDs, metrics, traces, dashboards and alert thresholds.',
    'Complete a security review covering JWT claims, secret rotation, OAuth redirect/cookie settings, provider signatures and dependency vulnerabilities.',
    'Run repeatable load tests and publish latency, throughput, error and resource-use results.'
])
add_heading('6.2.2 Product Enhancements',3)
add_bullets([
    'Unified inbox and business messaging, supported by the optional messaging service.',
    'Firebase push notifications for low stock, follow-up dates, payment state and expiring reservations.',
    'Refresh-token rotation, device/session management and user-initiated revocation.',
    'Offline-first mobile capture with conflict-aware synchronisation for unreliable connectivity.',
    'Multi-branch stock, transfer orders, role permissions, purchase orders and richer supplier analytics.',
    'Tax-aware receipts, Nepali calendar/localisation options and production payment-provider certification.',
    'Forecast evaluation using explainable historical models and explicit confidence ranges.',
    'Audit exports, backup/restore processes and configurable data-retention policies.'
])
add_heading('6.2.3 Research and Evaluation Enhancements',3)
add_body('Future study should involve a structured pilot with Nepalese small-business owners. Tasks should measure completion time, error rate, usability, perceived trust in AI suggestions and the effect of alerts on stock availability. Comparative evaluation against spreadsheets or an existing POS would provide stronger evidence of business value than developer testing alone.')

add_heading('6.3 Final Reflection',2)
add_body('The project combined software engineering, distributed data management, mobile and web design, security, testing and AI integration within a constrained academic schedule. The most important lesson was to make system states explicit. When identity, payment, reservation, cache and AI-review states are visible in the design, failure behavior can be reasoned about and tested rather than assumed.')

add_heading('REFERENCES',1)
refs = [
    'Docker, “Docker Compose documentation,” Docker Docs. Available: https://docs.docker.com/compose/ (accessed 3 August 2026).',
    'Expo, “Expo SDK reference,” Expo Documentation. Available: https://docs.expo.dev/versions/latest/ (accessed 3 August 2026).',
    'Expo, “Tutorial: Using React Native and Expo,” Expo Documentation. Available: https://docs.expo.dev/tutorial/introduction/ (accessed 3 August 2026).',
    'Flyway, “Flyway documentation,” Redgate Documentation. Available: https://documentation.red-gate.com/flyway (accessed 3 August 2026).',
    'Google, “Gemini API documentation,” Google AI for Developers. Available: https://ai.google.dev/gemini-api/docs (accessed 3 August 2026).',
    'Microsoft, “TypeScript documentation,” Available: https://www.typescriptlang.org/docs/ (accessed 3 August 2026).',
    'Neon, “Neon documentation,” Available: https://neon.com/docs (accessed 3 August 2026).',
    'Next.js, “Next.js Documentation,” Available: https://nextjs.org/docs (accessed 3 August 2026).',
    'OWASP Foundation, “REST Security Cheat Sheet,” OWASP Cheat Sheet Series. Available: https://cheatsheetseries.owasp.org/cheatsheets/REST_Security_Cheat_Sheet.html (accessed 3 August 2026).',
    'PostgreSQL Global Development Group, “PostgreSQL 16 Documentation: Transactions,” Available: https://www.postgresql.org/docs/16/tutorial-transactions.html (accessed 3 August 2026).',
    'Redis, “Redis documentation,” Available: https://redis.io/docs/latest/ (accessed 3 August 2026).',
    'Spring, “Spring Boot Reference Documentation,” Available: https://docs.spring.io/spring-boot/reference/ (accessed 3 August 2026).',
    'Spring, “Spring Cloud Gateway Reference Documentation,” Available: https://docs.spring.io/spring-cloud-gateway/reference/ (accessed 3 August 2026).',
    'Spring, “Spring Cloud Netflix / Eureka Reference,” Available: https://docs.spring.io/spring-cloud-netflix/reference/ (accessed 3 August 2026).',
    'SmartBiz project repository, source code, migrations, Docker configuration, tests, PROJECT_CONTEXT.md and PROGRESS.md, local academic project, 2026.'
]
for i, ref in enumerate(refs, 1):
    p=doc.add_paragraph(style='SB Reference')
    p.paragraph_format.left_indent=Inches(.28); p.paragraph_format.first_line_indent=Inches(-.28)
    p.add_run(f'[{i}] {ref}')

add_heading('APPENDIX A: API SUMMARY',1)
add_body('The following table summarises the principal public route groups. Detailed request and response fields remain defined in source DTOs so that the appendix does not reproduce secrets or environment-specific values.')
add_table('Table A.1: Principal API groups', ['Route group','Representative operations'], [
    ['/auth','Signup, login, verify/resend email, forgot/reset password, Google OAuth and profile.'],
    ['/billing','Plans, subscription status, checkout, payments and provider callbacks/webhooks.'],
    ['/inventory/categories','Category CRUD.'],
    ['/inventory/products','Paginated CRUD, barcode, low stock, adjust/restock and image operations.'],
    ['/inventory/suppliers','CRUD, summary, ledger, payments and adjustments.'],
    ['/inventory/reservations','Create, commit and release stock reservations.'],
    ['/customers','Paginated CRUD, purchase totals, due amounts and interactions.'],
    ['/leads','Paginated CRUD, stage updates and conversion.'],
    ['/sales','Create/import/list/get and today/weekly/trend analytics.'],
    ['/sales/payments','eSewa settings, create/check/cancel and callback operations.'],
    ['/ai','Query, insights, insight cards, invoice/voice/file parsing and import sessions.'],
], [2.0,4.25],9.1)

add_heading('APPENDIX B: CONFIGURATION CHECKLIST',1)
add_body('Configuration values are supplied through environment variables. Real values are intentionally omitted.')
add_table('Table B.1: Configuration categories', ['Category','Examples','Handling'], [
    ['Database','AUTH_DB_URL, INVENTORY_DB_URL, CRM_DB_URL, SALES_DB_URL, AI_DB_URL','Use service-specific database/schema; never commit passwords.'],
    ['Identity','JWT_SECRET, OAuth client ID/secret, verification URL','Rotate secrets and restrict redirect URLs.'],
    ['Discovery','EUREKA_URL','Use internal service URL in containers.'],
    ['AI','GEMINI_API_KEY, model URL','Keep key server-side; handle quota and model availability.'],
    ['Payments','eSewa/Stripe keys, callback URLs, enabled flags','Use UAT/test values until provider approval and production review.'],
    ['Clients','EXPO_PUBLIC_API_URL, web gateway URL','Point to reachable HTTPS gateway for deployment.'],
], [1.25,2.4,2.6],8.9)

add_heading('APPENDIX C: TEST EVIDENCE NOTE',1)
add_body_block('''Authoritative automated-test run environment: Java 21 with the repository Maven projects. Total discovered @Test methods: 58. Passed: 57. Failed: 1. The failing test is InsightServiceTest.buildInsightCards_returnsRestockSlowStockBundleAndCashFlowSignals in the AI service. Actual card types were SLOW_MOVING_STOCK, BUNDLE_OPPORTUNITY and CASH_FLOW_WARNING; RESTOCK_SOON was expected but absent.

This appendix records the observed result so that future work can reproduce the issue, decide whether the fixture or selection logic is incorrect, implement the change and rerun the complete reactor. No Java 25 tooling failure is counted as an application-test result because Java 21 is the project runtime.''')

# Global section and compatibility settings.
for s in doc.sections:
    configure_section(s)

settings = doc.settings.element
update = settings.find(qn('w:updateFields'))
if update is None:
    update = OxmlElement('w:updateFields')
    settings.append(update)
update.set(qn('w:val'),'true')

doc.core_properties.title = 'SmartBiz Final Year Project Report'
doc.core_properties.subject = 'Mobile-First Business Management System for Small Businesses in Nepal'
doc.core_properties.author = 'Sanskar Kayastha'
doc.core_properties.keywords = 'SmartBiz, Spring Boot, PostgreSQL, React Native, microservices, Nepal'
doc.core_properties.comments = 'Reconstructed from the supplied FYP report template; SmartBiz content based on the project repository.'

doc.save(OUT)
print(OUT)
