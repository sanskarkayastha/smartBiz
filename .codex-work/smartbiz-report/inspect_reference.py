from __future__ import annotations

import hashlib
import json
import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def length_value(value):
    if value is None:
        return None
    return {
        "emu": int(value),
        "inches": round(value.inches, 4),
        "points": round(value.pt, 2),
    }


def font_name(run):
    if run.font.name:
        return run.font.name
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        return rpr.rFonts.get(qn("w:ascii")) or rpr.rFonts.get(qn("w:hAnsi"))
    return None


def paragraph_record(index, paragraph, part="word/document.xml"):
    fmt = paragraph.paragraph_format
    runs = []
    for run in paragraph.runs:
        if not run.text:
            continue
        runs.append(
            {
                "text": run.text,
                "font": font_name(run),
                "size_pt": round(run.font.size.pt, 2) if run.font.size else None,
                "bold": run.bold,
                "italic": run.italic,
                "underline": bool(run.underline) if run.underline is not None else None,
                "color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
                "all_caps": run.font.all_caps,
            }
        )
    return {
        "part": part,
        "index": index,
        "style": paragraph.style.name if paragraph.style else None,
        "text": paragraph.text,
        "alignment": int(paragraph.alignment) if paragraph.alignment is not None else None,
        "left_indent": length_value(fmt.left_indent),
        "right_indent": length_value(fmt.right_indent),
        "first_line_indent": length_value(fmt.first_line_indent),
        "space_before": length_value(fmt.space_before),
        "space_after": length_value(fmt.space_after),
        "line_spacing": str(fmt.line_spacing) if fmt.line_spacing is not None else None,
        "keep_with_next": fmt.keep_with_next,
        "keep_together": fmt.keep_together,
        "page_break_before": fmt.page_break_before,
        "runs": runs,
    }


def main():
    source = Path(sys.argv[1]).resolve()
    output = Path(sys.argv[2]).resolve()
    doc = Document(source)
    result = {
        "reference": {
            "path": str(source),
            "sha256": hashlib.sha256(source.read_bytes()).hexdigest(),
            "size": source.stat().st_size,
        },
        "sections": [],
        "styles": [],
        "paragraphs": [],
        "tables": [],
        "headers": [],
        "footers": [],
        "package": [],
    }

    for index, section in enumerate(doc.sections, 1):
        result["sections"].append(
            {
                "index": index,
                "start_type": int(section.start_type),
                "orientation": int(section.orientation),
                "page_width": length_value(section.page_width),
                "page_height": length_value(section.page_height),
                "left_margin": length_value(section.left_margin),
                "right_margin": length_value(section.right_margin),
                "top_margin": length_value(section.top_margin),
                "bottom_margin": length_value(section.bottom_margin),
                "header_distance": length_value(section.header_distance),
                "footer_distance": length_value(section.footer_distance),
                "different_first_page": section.different_first_page_header_footer,
                "header_linked": section.header.is_linked_to_previous,
                "footer_linked": section.footer.is_linked_to_previous,
            }
        )

    for style in doc.styles:
        if style.type != 1:
            continue
        pf = style.paragraph_format
        font = style.font
        result["styles"].append(
            {
                "name": style.name,
                "base_style": style.base_style.name if style.base_style else None,
                "font": font.name,
                "size_pt": round(font.size.pt, 2) if font.size else None,
                "bold": font.bold,
                "italic": font.italic,
                "underline": bool(font.underline) if font.underline is not None else None,
                "color": str(font.color.rgb) if font.color and font.color.rgb else None,
                "alignment": int(pf.alignment) if pf.alignment is not None else None,
                "left_indent": length_value(pf.left_indent),
                "first_line_indent": length_value(pf.first_line_indent),
                "space_before": length_value(pf.space_before),
                "space_after": length_value(pf.space_after),
                "line_spacing": str(pf.line_spacing) if pf.line_spacing is not None else None,
                "keep_with_next": pf.keep_with_next,
                "page_break_before": pf.page_break_before,
            }
        )

    result["paragraphs"] = [paragraph_record(i, p) for i, p in enumerate(doc.paragraphs)]

    for table_index, table in enumerate(doc.tables):
        rows = []
        for row_index, row in enumerate(table.rows):
            cells = []
            for cell_index, cell in enumerate(row.cells):
                cells.append(
                    {
                        "index": cell_index,
                        "width": length_value(cell.width),
                        "vertical_alignment": int(cell.vertical_alignment) if cell.vertical_alignment is not None else None,
                        "paragraphs": [
                            paragraph_record(i, p, f"word/document.xml/table[{table_index}]/row[{row_index}]/cell[{cell_index}]")
                            for i, p in enumerate(cell.paragraphs)
                        ],
                    }
                )
            rows.append({"index": row_index, "cells": cells})
        result["tables"].append(
            {
                "index": table_index,
                "style": table.style.name if table.style else None,
                "rows": rows,
                "column_widths": [length_value(c.width) for c in table.rows[0].cells] if table.rows else [],
            }
        )

    seen_headers = set()
    seen_footers = set()
    for section_index, section in enumerate(doc.sections, 1):
        for kind, part, seen, target in (
            ("header", section.header, seen_headers, result["headers"]),
            ("footer", section.footer, seen_footers, result["footers"]),
        ):
            part_name = str(part.part.partname)
            if part_name in seen:
                continue
            seen.add(part_name)
            target.append(
                {
                    "section": section_index,
                    "part": part_name,
                    "paragraphs": [paragraph_record(i, p, part_name) for i, p in enumerate(part.paragraphs)],
                }
            )

    with zipfile.ZipFile(source) as package:
        for info in sorted(package.infolist(), key=lambda x: x.filename):
            data = package.read(info.filename)
            result["package"].append(
                {
                    "path": info.filename,
                    "size": info.file_size,
                    "sha256": hashlib.sha256(data).hexdigest(),
                }
            )

    role_counter = Counter(p["style"] for p in result["paragraphs"])
    result["summary"] = {
        "paragraph_count": len(result["paragraphs"]),
        "table_count": len(result["tables"]),
        "style_counts": dict(role_counter.most_common()),
    }

    output.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")


if __name__ == "__main__":
    main()
